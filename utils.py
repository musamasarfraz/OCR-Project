from docx import Document
from io import BytesIO

def create_docx(text_blocks: list) -> BytesIO:
    doc = Document()
    doc.add_heading('Extracted OCR Text', 0)

    for i, text in enumerate(text_blocks):
        doc.add_heading(f'File {i+1}', level=1)
        doc.add_paragraph(text)
        doc.add_page_break()

    # Save to a byte stream instead of a physical file
    target_stream = BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    return target_stream